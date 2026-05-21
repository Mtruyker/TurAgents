from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path(r"C:\Users\kab\Desktop\TourAgency1\Kursovaya_TourAgency.docx")
SCREENS = Path(r"C:\Users\kab\Desktop\TourAgency1\coursework_screenshots")

ITEMS = [
    ("01_tours.png", "Рисунок Г.1 – Экранная форма подбора путевок"),
    ("02_clients.png", "Рисунок Г.2 – Экранная форма работы с клиентами"),
    ("03_bookings_payments.png", "Рисунок Г.3 – Экранная форма учета заявок и оплат"),
    ("04_contract.png", "Рисунок Г.4 – Экранная форма формирования договора"),
    ("05_references.png", "Рисунок Г.5 – Экранная форма справочников направлений и отелей"),
]


def remove_tail_from(doc: Document, marker: str) -> None:
    start = None
    for i, paragraph in enumerate(doc.paragraphs):
        if marker in paragraph.text:
            start = i
            break
    if start is None:
        return
    for paragraph in list(doc.paragraphs[start:]):
        element = paragraph._element
        element.getparent().remove(element)


def set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str):
    doc.add_page_break()
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 2"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True)


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_screenshots(doc: Document):
    add_heading(doc, "Приложение Г. Экранные формы программного модуля")
    for index, (filename, caption) in enumerate(ITEMS):
        if index:
            doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.add_run().add_picture(str(SCREENS / filename), width=Cm(15.8))
        add_caption(doc, caption)


def main():
    doc = Document(DOCX)
    remove_tail_from(doc, "?????????? ?.")
    remove_tail_from(doc, "Приложение Г.")
    add_screenshots(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
